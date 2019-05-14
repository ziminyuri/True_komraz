from django.shortcuts import render, get_object_or_404, redirect
from .models import *
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from django.http import HttpResponse
import os
import re
import hashlib

#Костыльная форма авторизации
#Так делать не нужно!

# Функция получения хеша
def get_hash(password):
    hash = hashlib.sha256()
    hash.update(password.encode('utf-8'))
    password_hash = hash.hexdigest()
    return password_hash

# Session of users
user_master = "0"
user_sales = "0"
user_admin = "0"

def login(request):
    return render(request, 'top_komraz/login.html')


# Create your views here.
@csrf_exempt
def login(request):
    el = " "  # Ошибки на форму
    global user_admin
    global user_master
    global user_sales
    user_admin = "0"
    user_master = "0"
    user_sales = "0"

    if request.method == 'POST':
        print("Зашли")
        re_data = request.POST
        try:
            username = re_data['username']
            passw = re_data['pas']
            password=get_hash(passw)
            temp = get_object_or_404(Login1, login=username,password = password)

            #Мастер
            if temp.Id_Master != None:
                user_master = temp.Id_Master
                return redirect('master')

            #Cотрудник отдела по работе с обращениями
            if temp.Id_Employee != None:
                user_sales = temp.Id_Employee
                return redirect('index')

            # Администратор
            if (temp.Id_Employee ==None) and (temp.Id_Master ==None):
                    user_admin = '1'
                    return redirect('adminpage')

        except Exception as e:
            print("Поподаем в exception")
            print(e)
            el = "Неверная пара логин/пароль. Повторите попытку."

    return render(request, 'top_komraz/login.html',{'e': el})

def index(request):
    print(user_admin)
    print(user_sales)
    print(user_master)
    return render(request, 'top_komraz/main.html',{'master': user_master, 'sales':user_sales, 'admin':user_admin})

def master(request):
    return render(request, 'top_komraz/master.html',{'master': user_master, 'sales':user_sales, 'admin':user_admin})

def adminpage(request):
    return render(request, 'top_komraz/adminpage.html',{'master': user_master, 'sales':user_sales, 'admin':user_admin})

##################################################
############### СПРАВОЧНИКИ #####################
#################################################

#Услуги
def show_kind_services(request):
    services = KindService.objects.all()
    return render(request, 'top_komraz/kind_services.html',{'services': services,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Запчасти
def show_kind_detail(request):
    detail = KindDetail.objects.all()
    return render(request, 'top_komraz/kind_detail.html',{'details': detail,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Устройства
def show_kind_device(request):
    device = KindDevice.objects.all()
    return render(request, 'top_komraz/kind_device.html',{'devices': device,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Производители
def show_kind_creator(request):
    creator = Creator.objects.all()
    return render(request, 'top_komraz/creator.html',{'creators': creator,'master': user_master, 'sales':user_sales, 'admin':user_admin})

##################################################
############### Документы ########################
#################################################

#Акт приема
def show_acceptance_certificate(request):
    orders = Orderentity_fix2.objects.all()
    return render(request, 'top_komraz/acceptance_certificate.html',{'orders': orders,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Заполненный акт приема
@csrf_exempt
def download_act_enter(request):
    document = Document('top_komraz/act_enter_template.docx')
    re_data = request.POST
    number_act = re_data['number_act']
    fio = re_data['fio']
    address = re_data['address']
    phone = re_data['phone']
    model = re_data['model']
    imei = re_data['imei']
    appearance = re_data['appearance']
    name_work = re_data['name_work']
    defect = re_data['defect']
    employee = re_data['employee']
    date = re_data['date']

    for p in document.paragraphs:
        p.text = re.sub(r'НОМЕР', number_act, p.text)
        p.text = re.sub(r'КЛИЕНТ', fio, p.text)
        p.text = re.sub(r'АДРЕС', address, p.text)
        p.text = re.sub(r'ТЕЛЕФОН', phone, p.text)
        p.text = re.sub(r'НАИМЕНОВАНИЕ', model, p.text)
        p.text = re.sub(r'МОДЕЛЬ', imei, p.text)
        p.text = re.sub(r'ОПИСАНИЕ', appearance, p.text)
        p.text = re.sub(r'ВИД РЕМОНТА', name_work, p.text)
        p.text = re.sub(r'ДЕФЕКТ', defect, p.text)
        p.text = re.sub(r'СОТРУДНИК', employee, p.text)
        p.text = re.sub(r'ДАТА', date, p.text)

    document.save('top_komraz/act_enter_downloads.docx')
    filename = 'top_komraz/act_enter_downloads.docx'
    data = open(filename, "rb").read()
    response = HttpResponse(data, content_type='application/docx')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Disposition'] = 'attachment; filename=act_enter_downloads.docx'

    return response

#Форма акта приема
@csrf_exempt
def download_act(request):
    document = Document('top_komraz/act_enter.docx')
    document.save('top_komraz/act_downloads.docx')
    filename = 'top_komraz/act_downloads.docx'
    data = open(filename, "rb").read()
    response = HttpResponse(data, content_type='application/docx')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Disposition'] = 'attachment; filename=act_downloads.docx'

    return response

#Договор на поставку
@csrf_exempt
def show_contract_supply(request):
    deliverys = Request_delivery3.objects.all()
    if request.method == 'POST':
        re_data = request.POST

    # Добавление
        if re_data['_method'] == "POST":
            detail = get_object_or_404(Detail_fix2, Id_Detail=re_data['id_detail'])
            provider = get_object_or_404(Provider, ID_Provider=re_data['id_provider'])
            Request_delivery3.objects.create(ID_Detail=detail, Quantity=re_data['id_quantity'],Date=re_data['id_date'],ID_Provider = provider )
            return redirect('show_contract_supply')

    return render(request, 'top_komraz/contract_supply.html',{'deliverys': deliverys,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Договор на поставку: Добавление
def show_contract_supply_add(request):
    provider = Provider.objects.all()
    detail = Detail_fix2.objects.all()
    return render(request, 'top_komraz/contract_supply_add.html',{'providers': provider, 'details': detail,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Заполненный договор на поставку
@csrf_exempt
def download_supply(request):
    document = Document('top_komraz/supply_template.docx')
    re_data = request.POST
    date1 = re_data['date']
    number_act = re_data['number_act']
    provider = re_data['provider']
    detail = re_data['detail']
    quantity = re_data['quantity']
    price = re_data['price']
    address = re_data['address']
    phone = re_data['phone']
    employee = re_data['employee']

    price1 = price.split(',')
    end_price = int(price1[0]) * int(quantity)
    price1 = str(end_price)

    for p in document.paragraphs:
        p.text = re.sub(r'НОМЕР', number_act, p.text)
        p.text = re.sub(r'ДАТА', date1, p.text)
        p.text = re.sub(r'ПОСТАВЩИК', provider, p.text)
        p.text = re.sub(r'ЧАСТИ', detail, p.text)
        p.text = re.sub(r'КОЛИЧЕСТВО', quantity, p.text)
        p.text = re.sub(r'ЦЕНА', price, p.text)
        p.text = re.sub(r'СТОИМОСТЬ', price1, p.text)
        p.text = re.sub(r'АДРЕСОК', address, p.text)
        p.text = re.sub(r'ТЕЛЕФОН', phone, p.text)
        p.text = re.sub(r'СОТРУДНИК', employee, p.text)

    document.save('top_komraz/download_supply.docx')
    filename = 'top_komraz/download_supply.docx'
    data = open(filename, "rb").read()
    response = HttpResponse(data, content_type='application/docx')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Disposition'] = 'attachment; filename=download_supply.docx'

    return response

#Форма договора на поставку
@csrf_exempt
def download_supply_form(request):
    document = Document('top_komraz/supply.docx')
    document.save('top_komraz/downloads_supply_form.docx')
    filename = 'top_komraz/downloads_supply_form.docx'
    data = open(filename, "rb").read()
    response = HttpResponse(data, content_type='application/docx')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Disposition'] = 'attachment; filename=downloads_supply_form.docx'

    return response

#Акт выполненных работ
def show_certificate_of_completion(request):
    certificate = Certificate_fix2.objects.all()
    return render(request, 'top_komraz/certificate_of_completion.html',{'certificates': certificate,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Заполненный акт выполненных работ
@csrf_exempt
def download_act_end_template(request):
    document = Document('top_komraz/act_end_template.docx')
    re_data = request.POST

    number_act = re_data['number_act']
    fio = re_data['fio']
    address = re_data['address']
    phone = re_data['phone']
    model = re_data['model']
    imei = re_data['imei']
    name_work = re_data['name_work']
    defect = re_data['defect']
    employee = re_data['employee']
    date = re_data['date']

    for p in document.paragraphs:
        p.text = re.sub(r'НОМЕР', number_act, p.text)
        p.text = re.sub(r'СОТРУДНИК', employee, p.text)
        p.text = re.sub(r'ДАТА', date, p.text)
        p.text = re.sub(r'КЛИЕНТ', fio, p.text)
        p.text = re.sub(r'АДРЕС', address, p.text)
        p.text = re.sub(r'ТЕЛЕФОН', phone, p.text)
        p.text = re.sub(r'НАИМЕНОВАНИЕ', model, p.text)
        p.text = re.sub(r'МОДЕЛЬ', imei, p.text)
        p.text = re.sub(r'УСЛУГА', name_work, p.text)
        p.text = re.sub(r'ЧАСТЬ', defect, p.text)

    document.save('top_komraz/act_end_downloads.docx')
    filename = 'top_komraz/act_end_downloads.docx'
    data = open(filename, "rb").read()
    response = HttpResponse(data, content_type='application/docx')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Disposition'] = 'attachment; filename=act_end_downloads.docx'

    return response

#Форма акта выполненных работ
@csrf_exempt
def download_act_end(request):
    document = Document('top_komraz/act_end.docx')
    document.save('top_komraz/act_end_downloads.docx')
    filename = 'top_komraz/act_end_downloads.docx'
    data = open(filename, "rb").read()
    response = HttpResponse(data, content_type='application/docx')
    response['Content-Length'] = os.path.getsize(filename)
    response['Content-Disposition'] = 'attachment; filename=act_end_downloads.docx'

    return response

##################################################
############### УЧЕТ  ############################
#################################################


#Обращения
@csrf_exempt
def show_appeals(request):
    orders = Orderentity_fix2.objects.all()
    if request.method == 'POST':
        re_data = request.POST

        #Добавление
        if re_data['_method'] == "POST":
            #Устройство
            kind_device = get_object_or_404(KindDevice, Id_KindDevice=re_data['id_kind'])
            creator = get_object_or_404(Creator, Id_creator=re_data['id_creator'])
            Device.objects.create(Model_device= re_data['id_model'], IMEI=re_data['id_imei'], Id_KindDevice=kind_device, Id_creator=creator)


            #Акт приема
            customer = get_object_or_404(Customer, Id_customer=re_data['id_customer'])
            employee = get_object_or_404(Employee, Id_Employee= 1)      #Вот эту строку надо пофиксить
            device = Device.objects.last()
            service = get_object_or_404(Service, IdService=re_data['id_service'])
            detail = get_object_or_404(Detail_fix2, Id_Detail=re_data['id_detail'])
            Contract_fix2.objects.create(ConclusionDate=re_data['id_date'], PeriodExecute=re_data['id_date_end'],
                                         Defect=re_data['id_defect'], Appearance=re_data['id_appearance'],
                                         Id_customer=customer, Id_Employee=employee, Id_device=device, ID_service=service,
                                         ID_Detail_fix=detail)

            #Заказ
            status = get_object_or_404(Status, Id_Status=re_data['id_status'])
            contract = Contract_fix2.objects.last()
            master = get_object_or_404(Master, Id_Master=re_data['id_master'])
            Orderentity_fix2.objects.create(Id_Status=status, Id_contract=contract, Id_Master=master)
            return redirect('show_appeals')

        # Редактирование
        if re_data['_method'] == "PUT":
            # Устройство
            kind_device = get_object_or_404(KindDevice, Id_KindDevice=re_data['id_kind'])
            creator = get_object_or_404(Creator, Id_creator=re_data['id_creator'])
            Device.objects.filter(Id_device=re_data['device']).update(Model_device=re_data['id_model'],
                                                                      IMEI=re_data['id_imei'], Id_KindDevice=kind_device,
                                                                        Id_creator=creator)

            # Акт приема
            customer = get_object_or_404(Customer, Id_customer=re_data['id_customer'])
            employee = get_object_or_404(Employee, Id_Employee=1)  # Вот эту строку надо пофиксить
            device = get_object_or_404(Device, Id_device=re_data['device'])
            service = get_object_or_404(Service, IdService=re_data['id_service'])
            detail = get_object_or_404(Detail_fix2, Id_Detail=re_data['id_detail'])
            Contract_fix2.objects.filter(Id_contract=re_data['contract']).update(ConclusionDate=re_data['id_date'],
                                                                                 PeriodExecute=re_data['id_date_end'],
                                                                                 Defect=re_data['id_defect'],
                                                                                 Appearance=re_data['id_appearance'],
                                                                                 Id_customer=customer,
                                                                                 Id_Employee=employee,
                                                                                 Id_device=device,
                                                                                 ID_service=service,
                                                                                 ID_Detail_fix=detail)

            # Заказ
            status = get_object_or_404(Status, Id_Status=re_data['id_status'])
            contract = Contract_fix2.objects.last()
            master = get_object_or_404(Master, Id_Master=re_data['id_master'])
            Orderentity_fix2.objects.filter(ID_Order=re_data['order']).update(Id_Status=status,Id_contract=contract, Id_Master=master)

        # Удаление
        if re_data['_method'] == "DELETE":

            order = re_data['order']
            print(order)
            data = Orderentity_fix2.objects.get(ID_Order= order)
            data.delete()

            data = Contract_fix2.objects.get(Id_contract=re_data['contract'])
            data.delete()

            data = Device.objects.get(Id_device=re_data['device'])
            data.delete()

            return redirect('show_appeals')



    return render(request, 'top_komraz/appeals.html',{'orders': orders,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Обращения: Добавиление
def show_appeals_add(request):
    orders = Orderentity_fix2.objects.all()
    statuss = Status.objects.all()
    masters = Master.objects.all()
    contracts = Contract_fix2.objects.all()
    devices = Detail_fix2.objects.all()
    creators = Creator.objects.all()
    kinddevices = KindDevice.objects.all()
    services = Service.objects.all()
    customres = Customer.objects.all()
    details = Detail_fix2.objects.all()
    return render(request, 'top_komraz/appeals_add.html',{'orders': orders, 'statuss': statuss, 'masters': masters, 'contracts': contracts, 'devices': devices,
                                                      'creators': creators, 'kinddevices': kinddevices, 'services': services, 'customres': customres,
                                                      'details': details,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Обращения: Редактирование
def show_appeals_update(request,id):
    orders = get_object_or_404(Orderentity_fix2, ID_Order=id)
    statuss = Status.objects.all()
    masters = Master.objects.all()
    contracts = Contract_fix2.objects.all()
    devices = Detail_fix2.objects.all()
    creators = Creator.objects.all()
    kinddevices = KindDevice.objects.all()
    services = Service.objects.all()
    customres = Customer.objects.all()
    details = Detail_fix2.objects.all()
    return render(request, 'top_komraz/appeals_update.html',{'orders': orders, 'statuss': statuss, 'masters': masters, 'contracts': contracts, 'devices': devices,
                                                      'creators': creators, 'kinddevices': kinddevices, 'services': services, 'customres': customres,
                                                      'details': details,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Запчасти в наличии
@csrf_exempt
def show_parts_stock(request):
    details = Detail_fix2.objects.all()
    if request.method == 'POST':
        re_data = request.POST

        # Добавление
        if re_data['_method'] == "POST":
            detail = get_object_or_404(KindDetail, Id_KindDetail=re_data['id_kind'])
            Detail_fix2.objects.create(NameDetail = re_data['id_name'], PriceDetail = re_data['id_price'],
                                       Market_price = re_data['id_market'],Number = re_data['id_quantity'],
                                       Id_KindDetail = detail )
            return redirect('show_parts_stock')

        #Редактирование
        if re_data['_method'] == "PUT":
            Detail_fix2.objects.filter(Id_Detail =re_data['detail']).update(Number=re_data['id_number'])
            return redirect('show_parts_stock')

    return render(request, 'top_komraz/parts_stock.html',{'details': details,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Запчасти в наличии: Добавиление
def show_parts_stock_add(request):
    details = Detail_fix2.objects.all()
    kind = KindDetail.objects.all()
    return render(request, 'top_komraz/parts_stock_add.html',{ 'details': details, 'kinds' : kind,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Запчасти в наличии: Редактирование количества
def show_parts_stock_update(request,id):
    detail = get_object_or_404(Detail_fix2,  Id_Detail =id)
    return render(request, 'top_komraz/parts_stock_update.html',{'detail': detail,'master': user_master, 'sales':user_sales, 'admin':user_admin})


#Запчасти реализованные
@csrf_exempt
def show_parts_sold(request):
    details = Implemented.objects.all()
    if request.method == 'POST':
        re_data = request.POST

        # Добавление
        if re_data['_method'] == "POST":
            detail = get_object_or_404(Detail_fix2, Id_Detail=re_data['id_detail'])
            Implemented.objects.create(ID_Detail=detail, Quantity=re_data['id_quantity'])
            return redirect('show_parts_sold')

        # Редактирование
        if re_data['_method'] == "PUT":
            Implemented.objects.filter(ID_certificate=re_data['detail']).update(Quantity=re_data['id_number'])
            return redirect('show_parts_sold')

        # Удаление
        if re_data['_method'] == "DELETE":
            data = Implemented.objects.get(ID_certificate=re_data['device'])
            data.delete()

            return redirect('show_parts_sold')
    return render(request, 'top_komraz/parts_sold.html',{'details': details,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Запчасти реализованные: Добавиление
def show_parts_sold_add(request):
    details = Implemented.objects.all()
    return render(request, 'top_komraz/parts_sold_add.html',{ 'details': details,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Запчасти реализованные: Редактирование количества
def show_parts_sold_update(request,id):
    detail = get_object_or_404(Implemented,  ID_certificate =id)
    return render(request, 'top_komraz/parts_sold_update.html',{'detail': detail,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Платежи
@csrf_exempt
def show_payments(request):
    orders = Orderentity_fix2.objects.all()
    if request.method == 'POST':
        re_data = request.POST

        # Редактирование
        if re_data['_method'] == "PUT":
            status = get_object_or_404(Status, NameStatus="В очереди")
            print(status)
            Orderentity_fix2.objects.filter(ID_Order=re_data['order']).update(Id_Status=status)
            return redirect('show_payments')



    return render(request, 'top_komraz/Payments.html',{'orders': orders,'master': user_master, 'sales':user_sales, 'admin':user_admin})

##################################################
############### ПЛАНЫ ############################
#################################################

#Закупок
@csrf_exempt
def show_procurement(request):
    plan = Plan.objects.all()
    if request.method == 'POST':
        re_data = request.POST

    # Добавление
        if re_data['_method'] == "POST":
            detail = get_object_or_404(Detail_fix2, Id_Detail=re_data['id_detail'])
            Plan.objects.create(ID_Detail=detail, Quantity=re_data['id_quantity'])
            return redirect('show_procurement')

            # Удаление
        if re_data['_method'] == "DELETE":
            data = Plan.objects.get(id_plan_delivery=re_data['device'])
            data.delete()
            return redirect('show_procurement')

    return render(request, 'top_komraz/Procurement.html',{'plans': plan,'master': user_master, 'sales':user_sales, 'admin':user_admin})

#Закупки: Добавление
def show_procurement_add(request):
    details = Detail_fix2.objects.all()
    return render(request, 'top_komraz/Procurement_add.html',{'details': details,'master': user_master, 'sales':user_sales, 'admin':user_admin})


